from docx import Document
from docx.shared import Inches
import pandas as pd
from datetime import datetime
import os

# Load the CSV summary
df = pd.read_csv(r"C:\Users\Hp\PycharmProjects\pythonProject2\data\data\discovery_candidates_summary.csv")

# Create the document
doc = Document()
doc.add_heading("🧭 Archaeological Discovery Candidate Report", 0)

# Add date
doc.add_paragraph(f"📅 Date: {datetime.now().strftime('%B %d, %Y')}")

# Summary count
doc.add_paragraph(f"🗺️ Total New Discovery Candidates: {len(df)}")

# Add a table
doc.add_heading("📊 Candidate Summary", level=1)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Area (sq km)'
hdr_cells[1].text = 'Centroid Latitude'
hdr_cells[2].text = 'Centroid Longitude'

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = f"{row['area_sqkm']:.2f}"
    row_cells[1].text = f"{row['centroid_lat']:.6f}"
    row_cells[2].text = f"{row['centroid_lon']:.6f}"

# Optional: Embed map image if it exists
image_path = "data/discovery_candidates_summary.png"
if os.path.exists(image_path):
    doc.add_page_break()
    doc.add_heading("🗺️ Map Overview", level=1)
    doc.add_picture(image_path, width=Inches(5.5))
else:
    doc.add_paragraph("🖼️ Map image not found. Skipping map inclusion.")

# Save the document
output_path = "data/discovery_candidate_report.docx"
doc.save(output_path)
print(f"✅ Word report saved to: {output_path}")
